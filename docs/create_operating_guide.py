from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "docs" / "MCA SaaS Hub License Manager Operating Guide.docx"


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run)
    return para


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, size=9, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            run = cells[index].paragraphs[0].add_run("" if value is None else str(value))
            set_run_font(run, size=9)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        set_run_font(run)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.color.rgb = RGBColor(11, 27, 38)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MCA SaaS Hub License Manager")
    set_run_font(run, size=24, bold=True, color=(53, 37, 143))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Operating guide for customer-managed use")
    set_run_font(run, size=12, color=(85, 96, 109))

    add_para(doc, "Version: 0.1.0")
    add_para(doc, "Package folder: mca-saas-hub-license-manager")
    add_para(
        doc,
        "Disclaimer: This tool is independent work and is not a Microsoft product, service, endorsed solution, or supported offering. It is provided as-is for customer-managed use. Microsoft has no responsibility for this tool, its operation, or any billing or commerce changes made with it. Customers are responsible for validating behavior, permissions, billing impact, and all commerce changes before use.",
    )

    doc.add_heading("1. What this tool is", level=1)
    add_para(
        doc,
        "This tool helps operators manage active MCA-billed subscriptions from a local workstation. It has a browser UI for day-to-day use and a standalone PowerShell script for customers who prefer automation.",
    )
    add_para(
        doc,
        "It doesn't keep a database. It doesn't store credentials. It uses the signed-in Azure CLI user and calls Microsoft Billing and Microsoft SaaS Hub APIs through Azure Resource Manager.",
    )
    add_table(
        doc,
        ["Area", "What it does"],
        [
            ["Inventory", "Loads active products and subscriptions across billing accounts the signed-in user can access. Exports the result to CSV."],
            ["Change seats", "Sets a new total seat quantity for an existing subscription. This can add or reduce seats when commerce policy allows it."],
            ["Recurring billing", "Turns recurring billing on or off for an existing subscription."],
            ["Move tenant", "Moves an existing subscription to another active associated tenant."],
            ["Operations", "Polls long-running SaaS Hub operations until a final state is returned."],
        ],
    )

    doc.add_heading("2. What is not included", level=1)
    add_para(
        doc,
        "New product purchase isn't exposed in this handover build. During testing it wasn't reliable enough across billing profiles. Product search and term resolution can fail depending on billing group, market, or product state.",
    )
    add_para(
        doc,
        "Change billing plan by product code is also not exposed. The trace showed the low-level API shape, but a normal user won't know the target product code.",
    )

    doc.add_heading("3. Installation", level=1)
    add_para(
        doc,
        "Unzip or copy the mca-saas-hub-license-manager folder to the operator workstation. Open PowerShell in that folder and run the installer. The installer checks for Azure CLI, Node.js, and npm. If winget is available, it can install missing prerequisites.",
    )
    add_table(
        doc,
        ["Step", "Command", "What it does"],
        [
            ["Install prerequisites and build", ".\\Install-McaSaasHub.ps1", "Checks Azure CLI, Node.js, and npm, then runs npm install and npm run build."],
            ["Sign in", "az login", "Signs the operator into Azure. Use the account that has MCA billing roles."],
            ["Start the UI", ".\\Start-McaSaasHub.ps1 -OpenBrowser", "Starts the local web app and opens http://localhost:3333."],
            ["Script-only inventory export", ".\\scripts\\McaSaasHub.ps1 -Action ExportInventory -OutputPath .\\inventory.csv", "Exports active inventory without using the UI."],
        ],
    )
    add_para(
        doc,
        "If the customer doesn't allow winget installs, install Azure CLI and Node.js LTS manually, close and reopen PowerShell, then run .\\Install-McaSaasHub.ps1 -SkipWingetInstall.",
    )
    add_para(doc, "Manual setup is also supported. Use this when software installation is managed by desktop engineering or when winget is blocked.")
    add_table(
        doc,
        ["Step", "Command"],
        [
            ["Confirm Azure CLI", "az --version"],
            ["Confirm Node.js", "node --version"],
            ["Confirm npm", "npm --version"],
            ["Sign in", "az login"],
            ["Install app packages", "npm install"],
            ["Build the UI and local API", "npm run build"],
            ["Start the local app", "npm start"],
            ["Open the app", "http://localhost:3333"],
        ],
    )

    doc.add_heading("4. Architecture", level=1)
    add_table(
        doc,
        ["Component", "Location", "Notes"],
        [
            ["Browser UI", "src/client", "React and Vite. Runs locally in a browser at http://localhost:3333 by default."],
            ["Local API", "src/server", "Node and Express. Calls Azure Resource Manager using the Azure CLI user token."],
            ["PowerShell script", "scripts/McaSaasHub.ps1", "Standalone script. It doesn't need Node, npm, Scout, or the UI."],
            ["Docs", "README.md and docs folder", "Customer setup notes and this operating guide."],
        ],
    )

    doc.add_heading("5. Authentication model", level=1)
    add_para(
        doc,
        "The UI and script both rely on Azure CLI sign-in. The user runs az login, or uses the UI sign-in button, then the local API gets an ARM token with az account get-access-token. The browser doesn't receive the token.",
    )
    add_table(
        doc,
        ["Step", "Mechanism"],
        [
            ["Sign in", "az login or device-code sign-in from the UI"],
            ["Token", "az account get-access-token --resource https://management.azure.com"],
            ["Authorization", "Microsoft Entra permissions and MCA billing roles decide what the user can see or change"],
            ["Storage", "No credential store is created by the tool"],
        ],
    )

    doc.add_heading("6. UI operating guide", level=1)
    add_table(
        doc,
        ["UI section", "How to use it"],
        [
            ["Sign in", "Check the current account and tenant. Use Sign in / switch account when the wrong account is active."],
            ["Inventory view", "Click Load inventory. Review active products across accessible MCA billing accounts. Use Download CSV or Pop out inventory when needed."],
            ["Billing scope", "Load billing access, then choose the billing account, billing profile, and invoice section. This scope is used by subscription management actions."],
            ["Change seats", "Load existing subscriptions, choose a subscription, enter the new total quantity, then click Change seats."],
            ["Recurring billing", "Choose the subscription, pick On or Off, then click Update recurring billing."],
            ["Move subscription tenant", "Choose the subscription, select an active associated tenant, then click Move selected subscription."],
            ["Operations", "Use Poll status for accepted async operations. Wait for Succeeded, Failed, or Canceled."],
        ],
    )

    doc.add_heading("7. PowerShell operating guide", level=1)
    add_para(doc, "Copy scripts/McaSaasHub.ps1 to a workstation that has Azure CLI installed. Run az login first.")
    add_table(
        doc,
        ["Task", "Example"],
        [
            ["Show active account", ".\\McaSaasHub.ps1 -Action WhoAmI"],
            ["Export inventory", ".\\McaSaasHub.ps1 -Action ExportInventory -OutputPath .\\inventory.csv"],
            ["List subscriptions", ".\\McaSaasHub.ps1 -Action ListBillingSubscriptions -BillingAccountName \"<billing-account-name>\""],
            ["Change seats", ".\\McaSaasHub.ps1 -Action ChangeSeats -BillingSubscriptionName \"<id>\" -Quantity 10 -Force"],
            ["Recurring billing off", ".\\McaSaasHub.ps1 -Action SetRecurringBilling -BillingSubscriptionName \"<id>\" -AutoRenew Off -Force"],
            ["Move tenant", ".\\McaSaasHub.ps1 -Action MoveSubscriptionTenant -BillingSubscriptionName \"<id>\" -ProvisioningTenantId \"<tenant-id>\" -Force"],
            ["Poll operation", ".\\McaSaasHub.ps1 -Action PollOperation -OperationUrl \"<operation-url>\""],
        ],
    )
    add_para(doc, "Mutating commands require -Force. That's intentional. It makes the operator pause and review the command before changing commerce state.")

    doc.add_heading("8. API map", level=1)
    add_table(
        doc,
        ["Capability", "API", "Method", "Notes"],
        [
            ["List billing accounts", "/providers/Microsoft.Billing/billingAccounts?api-version=2024-04-01", "GET", "Returns billing accounts visible to the signed-in user."],
            ["List billing profiles", "/providers/Microsoft.Billing/billingAccounts/{billingAccountName}/billingProfiles?api-version=2024-04-01", "GET", "Used for scope selection."],
            ["List invoice sections", "/providers/Microsoft.Billing/billingAccounts/{billingAccountName}/billingProfiles/{billingProfileName}/invoiceSections?api-version=2024-04-01", "GET", "Used for invoice section selection."],
            ["List billing subscriptions", "/providers/Microsoft.Billing/billingAccounts/{billingAccountName}/billingSubscriptions?api-version=2024-04-01", "GET", "Source for inventory and existing subscriptions. Tool keeps only Active rows."],
            ["List associated tenants", "/providers/Microsoft.Billing/billingAccounts/{billingAccountName}/associatedTenants?api-version=2020-11-01-privatepreview", "GET", "Used for move tenant target list."],
            ["Change seats", "/providers/Microsoft.SaaSHub/saasResources/{billingSubscriptionName}?api-version=2025-07-01-preview", "PATCH", "Body uses properties.quantity."],
            ["Recurring billing", "/providers/Microsoft.SaaSHub/saasResources/{billingSubscriptionName}?api-version=2025-07-01-preview", "PATCH", "Body uses properties.autoRenew set to On or Off."],
            ["Move tenant", "/providers/Microsoft.SaaSHub/saasResources/{billingSubscriptionName}?api-version=2025-07-01-preview", "PATCH", "Body uses properties.provisioningTenantId."],
            ["Poll operation", "Azure-AsyncOperation or Location URL returned by SaaS Hub", "GET", "Poll until a final status is returned."],
        ],
    )

    doc.add_heading("9. Inventory columns", level=1)
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["Azure / Non-Azure", "Derived classification. Reservations, Azure plans, and Azure-like resources are marked Azure. Seat-based Microsoft 365 style subscriptions are marked Non-Azure."],
            ["Billing Account and ID", "The billing account display name and system name."],
            ["Billing Profile and ID", "The billing profile display name and system name."],
            ["Invoice Section and ID", "The invoice section display name and system name."],
            ["Subscription ID", "The billing subscription name used by APIs."],
            ["Product Name", "Best available product label, usually SKU description or product type."],
            ["Quantity", "Seat or unit quantity when present."],
            ["Term", "Term duration such as P1M, P1Y, or P3Y."],
            ["Expiry Date", "Mapped from termEndDate."],
            ["Provisioning Tenant ID", "Tenant where the subscription is provisioned, when present."],
        ],
    )

    doc.add_heading("10. Safety and operating rules", level=1)
    add_bullets(
        doc,
        [
            "Run inventory first before making any changes.",
            "Use a test subscription for validation before changing production subscriptions.",
            "Review the selected subscription and current quantity before changing seats.",
            "Seat reductions can be restricted after the post-purchase adjustment window.",
            "Use recurring billing off when the goal is to stop renewal. Don't treat it as immediate cancellation.",
            "Save operation URLs from mutating actions and poll them to completion.",
            "Keep the script under source control if the customer changes it.",
        ],
    )

    doc.add_heading("11. Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "What to check"],
        [
            ["No billing accounts appear", "Wrong Azure CLI account or missing billing role", "Run az account show. Confirm MCA billing role assignments."],
            ["Inventory is empty", "No active billing subscriptions visible to the user", "Try ListBillingAccounts, then ListBillingSubscriptions for a known billing account."],
            ["Change seats fails", "Commerce policy or permission block", "Check quantity, post-purchase window, and billing role."],
            ["Recurring billing fails", "Subscription doesn't allow auto-renew change", "Check operation error and subscription policy."],
            ["Move tenant fails", "Target tenant isn't active or associated", "Run ListAssociatedTenants and use a tenant with provisioningState Active."],
            ["Network errors to ARM", "Proxy, VPN, firewall, or TLS interception", "Test az rest against management.azure.com."],
        ],
    )

    doc.add_heading("12. Day-one acceptance test", level=1)
    add_bullets(
        doc,
        [
            "Run .\\Install-McaSaasHub.ps1 and confirm it completes without errors.",
            "Run az login with an account that has the required MCA billing role.",
            "Start the UI with .\\Start-McaSaasHub.ps1 -OpenBrowser.",
            "Load inventory and confirm active products appear.",
            "Download the inventory CSV and confirm it contains Billing Account, Billing Profile, Invoice Section, Product Name, Quantity, Term, and Expiry Date.",
            "Run .\\scripts\\McaSaasHub.ps1 -Action ExportInventory -OutputPath .\\inventory.csv and confirm the script-only path works.",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Independent work. Not a Microsoft product or supported offering. Validate before commerce changes.")
    set_run_font(run, size=8, color=(85, 96, 109))

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
